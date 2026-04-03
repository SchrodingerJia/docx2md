import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shape import InlineShape
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.enum.text import WD_BREAK

class DocxHandler:
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    def __init__(self, file_path=None, image_dir="assets/images"):
        """
        初始化 DocxHandler
        :param file_path: 如果提供路径则加载现有文档，否则创建一个新文档
        """
        self.file_path = file_path
        if file_path and os.path.exists(file_path):
            self.doc = Document(file_path)
        else:
            self.doc = Document()
        self.image_dir = image_dir
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)

    # --- 文件读取与解析 ---

    def get_all_text(self):
        """获取文档中所有段落的纯文本"""
        return [para.text for para in self.doc.paragraphs]

    def get_tables_data(self):
        """获取文档中所有表格的数据"""
        all_tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            all_tables.append(table_data)
        return all_tables
    
    def get_full_details(self):
        content_details = []
        for element in self.doc.element.body:
            # 1. 处理段落 (包括文本、图片、超链接、列表)
            if isinstance(element, CT_P):
                para = Paragraph(element, self.doc)
                content_details.append(self._parse_paragraph(para))
            
            # 2. 处理表格
            elif isinstance(element, CT_Tbl):
                table = Table(element, self.doc)
                content_details.append(self._parse_table(table))
            
            # 3. 分节符/分页符 (在 XML 中通常是 pPr 里的 sectPr，这里简单处理)
            # 也可以通过判断段落中是否有渲染分页符来识别
        return content_details
    
    def _parse_paragraph(self, para):
        """解析段落，处理 Run, Hyperlink, Image 和 Math"""
        runs_data = []
        list_lvl = self._get_list_level(para)
        
        # 遍历段落的所有子元素，包括文本(r), 超链接(hyperlink)和公式(oMath/oMathPara)
        for child in para._element.xpath('w:r | w:hyperlink | m:oMath | m:oMathPara'):
            tag = child.tag
            
            # 1. 处理数学公式 (oMath: 行内, oMathPara: 块级)
            if tag.endswith('oMath') or tag.endswith('oMathPara'):
                latex_code = self._omml_to_latex(child)
                runs_data.append({
                    "type": "formula",
                    "text": latex_code,
                    "is_block": tag.endswith('oMathPara')
                })
            # 2. 处理超链接
            elif tag.endswith('hyperlink'):
                for r_node in child.xpath('w:r'):
                    run = Run(r_node, para)
                    run_dict = self._get_run_dict(run)
                    # 获取链接关系ID
                    rId = child.get(qn('r:id'))
                    if rId:
                        run_dict["link_url"] = para.part.rels[rId].target_ref
                    runs_data.append(run_dict)
            # 3. 处理普通文本或图片 (w:r)
            elif tag.endswith('r'):
                run = Run(child, para)
                image_info = self._extract_images_from_run(run)
                if image_info:
                    runs_data.append({"type": "image", "src": image_info})
                else:
                    runs_data.append(self._get_run_dict(run))
        return {
            "type": "paragraph",
            "style": para.style.name if para.style else "Default",
            "list_level": list_lvl,
            "runs": runs_data
        }
    
    def _omml_to_latex(self, omml_element):
        """
        简单的 OMML 到 LaTeX 的转换逻辑。
        注意：完全精确的转换需要复杂的 XSLT，这里实现基础的字符提取。
        建议生产环境下考虑结合第三方库如 latex2mathml 的逆向思路。
        """
        # 提取公式中所有的文本内容 (m:t 标签)
        # 在 Word XML 中，公式的字符存储在 m:t 节点中
        t_nodes = omml_element.xpath('.//m:t', namespaces={'m': self.MATH_NS})
        parts = [node.text for node in t_nodes if node.text]
        return "".join(parts)
    
    def _get_run_dict(self, run):
        """提取普通文本块格式"""
        color = run.font.color.rgb if run.font.color and run.font.color.rgb else "Default"
        return {
            "type": "text",
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "color": str(color),
            "font_size": run.font.size.pt if run.font.size else None
        }
    
    def _extract_images_from_run(self, run):
        """从 Run 中提取图片并保存"""
        # 检查 Run 的 XML 中是否有 drawing 标签
        img_lst = run._element.xpath('.//pic:pic')
        if not img_lst:
            return None
        
        for pic in img_lst:
            blip = pic.xpath('.//a:blip/@r:embed')[0]
            image_part = self.doc.part.related_parts[blip]
            image_filename = os.path.basename(image_part.partname)
            image_path = os.path.join(self.image_dir, image_filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_part.blob)
            return image_path
        return None
    
    def _get_list_level(self, para):
        """获取段落的列表层级"""
        pPr = para._element.pPr
        if pPr is not None and pPr.numPr is not None:
            ilvl = pPr.numPr.ilvl
            if ilvl is not None:
                return ilvl.val # 返回 0, 1, 2... 表示层级
        return None
    
    def _parse_table(self, table):
        """解析表格"""
        table_info = {
            "type": "table",
            "rows": []
        }
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_paras = []
                # 单元格内可能包含多个段落
                for p in cell.paragraphs:
                    cell_paras.append({
                        "type": "paragraph",
                        "text": p.text,
                        "style": p.style.name if p.style else "Default",
                        "runs": self._extract_run_info(p)
                    })
                row_data.append(cell_paras)
            table_info["rows"].append(row_data)
        return table_info
    
    def _extract_run_info(self, paragraph:Paragraph):
        """提取段落中每个 Run 的格式"""
        run_details = []
        for run in paragraph.runs:
            # 获取颜色
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else "Default"
            
            info = {
                "type": "text",
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "color": str(color)
            }
            run_details.append(info)
        return run_details

    # --- 特定内容提取 ---

    def find_paragraphs_with_keyword(self, keyword):
        """查找包含特定关键字的段落"""
        return [para.text for para in self.doc.paragraphs if keyword in para.text]

    # --- 局部修改 ---

    def replace_text(self, old_text, new_text):
        """
        全局替换文本 (保留基本格式)
        注意：python-docx 的替换受限于 'run' 的切分，此方法处理简单的文本替换
        """
        for para in self.doc.paragraphs:
            if old_text in para.text:
                # 这种方式会保留段落整体样式，但可能会打乱细碎的 Run 样式
                para.text = para.text.replace(old_text, new_text)
        
        # 处理表格中的文本
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)

    def update_cell(self, table_index, row, col, new_value):
        """修改指定表格中某个单元格的内容"""
        try:
            table = self.doc.tables[table_index]
            table.cell(row, col).text = str(new_value)
        except IndexError:
            print(f"Error: Table {table_index} or Cell ({row},{col}) not found.")

    # --- 内容创建 ---

    def add_heading(self, text, level=1):
        """添加标题"""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text, style=None, align='left'):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p

    def add_table(self, rows, cols, data=None):
        """
        添加表格并填充数据
        :param data: 二维列表
        """
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'  # 添加默认边框
        if data:
            for i, row_data in enumerate(data):
                for j, value in enumerate(row_data):
                    table.cell(i, j).text = str(value)
        return table

    def add_picture(self, image_path, width_inch=2.0):
        """插入图片"""
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width_inch))
        else:
            print(f"Warning: Image {image_path} not found.")

    # --- 保存 ---

    def save(self, save_path=None):
        """保存文档"""
        path = save_path if save_path else self.file_path
        if not path:
            raise ValueError("Save path must be provided.")
        self.doc.save(path)
        print(f"Document saved to {path}")

# --- 使用示例 ---
if __name__ == "__main__":
    path = r'D:\陈信嘉2025\论文\信息论\2023311A11-陈信嘉-信息论导论课程报告-主题2.docx'
    reader = DocxHandler(path)
    details = reader.get_full_details()
    for item in details:
        if item["type"] == "paragraph":
            print(f"--- 段落 (样式: {item['style']}) ---")
            for run in item["runs"]:
                # 过滤掉空字符串的 run
                if not run["text"].strip(): continue
                
                fmt = []
                if run["bold"]: fmt.append("加粗")
                if run["italic"]: fmt.append("斜体")
                if run["underline"]: fmt.append("下划线")
                size = f"{run['font_size']}pt" if run['font_size'] else "默认大小"
                
                print(f"内容: '{run['text']}' | 格式: [{', '.join(fmt)}] | 字体: {run['font_name']} | 大小: {size}")